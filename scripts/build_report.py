"""Build the Cloud Computing Lab Report (.docx) for Mukesh Pant.

Reads screenshots from `screenshots/lab-NN/` and writes a single Word document
to `report/Mukesh_Pant_Cloud_Computing_Lab_Report.docx`.

Run:
    python scripts/build_report.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "screenshots"
OUT = ROOT / "report" / "Mukesh_Pant_Cloud_Computing_Lab_Report.docx"

# ---------------------------------------------------------------------------
# Visual design tokens
# ---------------------------------------------------------------------------

COLOR_PRIMARY = RGBColor(0x1F, 0x38, 0x64)   # deep academic blue (Heading 1)
COLOR_SECONDARY = RGBColor(0x2E, 0x74, 0xB5) # mid blue (Heading 3, accents)
COLOR_GREY = RGBColor(0x59, 0x59, 0x59)      # captions, sub-text
COLOR_RULE = RGBColor(0xBF, 0xBF, 0xBF)      # horizontal rule under lab titles

FONT_BODY = "Calibri"
FONT_HEADING = "Calibri"

SIZE_BODY = 12
SIZE_LAB_TITLE = 22
SIZE_H1 = 14
SIZE_H3 = 12
SIZE_CAPTION = 10
SIZE_TOC_ENTRY = 12
SIZE_COVER_TITLE = 30
SIZE_COVER_BIG = 22
SIZE_COVER_NAME = 18


# ---------------------------------------------------------------------------
# Low-level helpers (XML / styling)
# ---------------------------------------------------------------------------

def _set_run(run, *, bold=False, italic=False, size=None, color=None, font_name=None):
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if font_name is not None:
        run.font.name = font_name


def _para_spacing(p, *, space_before=0, space_after=4, line=1.15, keep_with_next=False, keep_together=False, page_break_before=False):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.widow_control = True
    if keep_with_next:
        pf.keep_with_next = True
    if keep_together:
        pf.keep_together = True
    if page_break_before:
        pf.page_break_before = True


def _add_horizontal_rule(paragraph, color=COLOR_RULE, size=8):
    """Add a bottom border to a paragraph as a horizontal rule under it."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*color))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_page_field(run):
    """Insert {PAGE} field into a run for live page numbers."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Section / page setup
# ---------------------------------------------------------------------------

def configure_sections(doc):
    """Set Moderate margins and configure header/footer with page numbers."""
    section = doc.sections[0]
    # Word "Moderate": 1" top/bottom, 0.75" left/right
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    # different-first-page so cover page has no header/footer
    section.different_first_page_header_footer = True

    # Header (skipped on first page)
    header = section.header
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = h_para.add_run("Cloud Computing Lab Report  |  Mukesh Pant  |  Roll No. 29")
    _set_run(hr, italic=True, size=10, color=COLOR_GREY)
    _add_horizontal_rule(h_para, color=COLOR_RULE, size=4)

    # Footer (skipped on first page)
    footer = section.footer
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = f_para.add_run("Page ")
    _set_run(label, size=10, color=COLOR_GREY)
    page_run = f_para.add_run()
    _set_run(page_run, size=10, color=COLOR_GREY)
    _add_page_field(page_run)


def configure_default_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(SIZE_BODY)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _cover_centered(doc, text, *, bold=False, italic=False, size=None, color=None,
                    space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, space_before=space_before, space_after=space_after, line=1.2)
    r = p.add_run(text)
    _set_run(r, bold=bold, italic=italic, size=size, color=color)
    return p


def add_cover_page(doc):
    # Top spacer
    for _ in range(2):
        doc.add_paragraph()

    _cover_centered(doc, "FAR WESTERN UNIVERSITY",
                    bold=True, size=SIZE_COVER_TITLE, color=COLOR_PRIMARY,
                    space_after=8)
    _cover_centered(doc, "Faculty of Engineering", bold=True, size=18, color=COLOR_SECONDARY)
    p_school = _cover_centered(doc, "School of Engineering", italic=True, size=14, color=COLOR_GREY)
    _add_horizontal_rule(p_school, color=COLOR_PRIMARY, size=12)

    for _ in range(3):
        doc.add_paragraph()

    _cover_centered(doc, "A LAB REPORT", italic=True, size=14, color=COLOR_GREY,
                    space_after=4)
    _cover_centered(doc, "ON", italic=True, size=12, color=COLOR_GREY,
                    space_after=4)
    p_title = _cover_centered(doc, "CLOUD COMPUTING", bold=True,
                              size=SIZE_COVER_BIG, color=COLOR_PRIMARY,
                              space_after=2)
    _cover_centered(doc, "(Practical)", italic=True, size=14, color=COLOR_SECONDARY,
                    space_after=4)
    _add_horizontal_rule(p_title, color=COLOR_PRIMARY, size=8)

    for _ in range(3):
        doc.add_paragraph()

    # Submitted by block
    _cover_centered(doc, "Submitted by:", italic=True, size=12, color=COLOR_GREY,
                    space_after=4)
    _cover_centered(doc, "Mukesh Pant", bold=True, size=SIZE_COVER_NAME,
                    color=COLOR_PRIMARY, space_after=2)
    _cover_centered(doc, "Roll No. 29", bold=True, size=14, color=COLOR_SECONDARY,
                    space_after=2)
    _cover_centered(doc, "VIII Semester  |  B.E. Computer Engineering",
                    size=12, color=COLOR_GREY, space_after=4)

    for _ in range(2):
        doc.add_paragraph()

    # Submitted to block
    _cover_centered(doc, "Submitted to:", italic=True, size=12, color=COLOR_GREY,
                    space_after=4)
    _cover_centered(doc, "Er. Robinson Pujara", bold=True, size=14,
                    color=COLOR_PRIMARY, space_after=2)
    _cover_centered(doc, "Lecturer, School of Engineering, FWU",
                    size=12, color=COLOR_GREY, space_after=4)

    for _ in range(3):
        doc.add_paragraph()

    _cover_centered(doc, "Signature:  ____________________",
                    size=12, color=COLOR_GREY, space_after=2)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Table of contents
# ---------------------------------------------------------------------------

def add_toc(doc, labs):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, space_before=4, space_after=8, keep_with_next=True)
    r = p.add_run("TABLE OF CONTENTS")
    _set_run(r, bold=True, size=20, color=COLOR_PRIMARY)
    _add_horizontal_rule(p, color=COLOR_PRIMARY, size=8)
    doc.add_paragraph()

    for i, lab in enumerate(labs, 1):
        para = doc.add_paragraph()
        _para_spacing(para, space_before=2, space_after=4, line=1.2)
        run_n = para.add_run(f"Lab {i}.   ")
        _set_run(run_n, bold=True, size=SIZE_TOC_ENTRY, color=COLOR_PRIMARY)
        run_t = para.add_run(lab["title"])
        _set_run(run_t, size=SIZE_TOC_ENTRY, color=COLOR_GREY)


# ---------------------------------------------------------------------------
# Per-lab elements
# ---------------------------------------------------------------------------

def _add_h1(doc, text, *, keep_with_next=True):
    p = doc.add_paragraph()
    _para_spacing(p, space_before=10, space_after=4, line=1.2,
                  keep_with_next=keep_with_next)
    r = p.add_run(text)
    _set_run(r, bold=True, size=SIZE_H1, color=COLOR_PRIMARY)
    return p


def _add_h3(doc, text, *, keep_with_next=True):
    p = doc.add_paragraph()
    _para_spacing(p, space_before=6, space_after=2, line=1.2,
                  keep_with_next=keep_with_next)
    r = p.add_run(text)
    _set_run(r, bold=True, size=SIZE_H3, color=COLOR_SECONDARY)
    return p


def _add_body(doc, text, *, justify=True, keep_together=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _para_spacing(p, space_before=0, space_after=6, line=1.3,
                  keep_together=keep_together)
    r = p.add_run(text)
    _set_run(r, size=SIZE_BODY)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _para_spacing(p, space_before=0, space_after=2, line=1.2)
    if p.runs:
        for r in p.runs:
            _set_run(r, size=SIZE_BODY)
    # add the text as a run if style didn't already
    if not p.text.strip():
        r = p.add_run(text)
        _set_run(r, size=SIZE_BODY)
    else:
        # the style sometimes inserts placeholder; rewrite
        for r in p.runs:
            r.text = ""
        r = p.add_run(text)
        _set_run(r, size=SIZE_BODY)


def _add_services_table(doc, services):
    # Two-column compact table for AWS Services Used.
    n = len(services)
    rows = (n + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    for col in table.columns:
        col.width = Inches(3.2)

    for idx, svc in enumerate(services):
        r, c = divmod(idx, 2)
        cell = table.cell(r, c)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        para = cell.paragraphs[0]
        _para_spacing(para, space_before=2, space_after=2, line=1.15)
        run_b = para.add_run("•  ")
        _set_run(run_b, bold=True, size=SIZE_BODY, color=COLOR_SECONDARY)
        run_t = para.add_run(svc)
        _set_run(run_t, size=SIZE_BODY)
    # ensure layout has bottom spacer
    after = doc.add_paragraph()
    _para_spacing(after, space_before=0, space_after=4, line=1.0)


def _add_screenshot(doc, image_path, caption):
    if not Path(image_path).exists():
        warn = doc.add_paragraph()
        wr = warn.add_run(f"[MISSING SCREENSHOT: {image_path}]")
        _set_run(wr, italic=True, color=RGBColor(0xCC, 0x00, 0x00))
        return

    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(pic_para, space_before=4, space_after=2, line=1.0,
                  keep_with_next=True, keep_together=True)
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(5.6))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(cap, space_before=0, space_after=8, line=1.15,
                  keep_together=True)
    cr = cap.add_run(caption)
    _set_run(cr, italic=True, size=SIZE_CAPTION, color=COLOR_GREY)


def add_lab(doc, n, lab, *, is_first=False):
    title_p = doc.add_paragraph()
    _para_spacing(title_p, space_before=0, space_after=8, line=1.2,
                  keep_with_next=True,
                  page_break_before=not is_first)
    tr = title_p.add_run(f"Lab {n}: {lab['title']}")
    _set_run(tr, bold=True, size=SIZE_LAB_TITLE, color=COLOR_PRIMARY)
    _add_horizontal_rule(title_p, color=COLOR_PRIMARY, size=12)

    # Objective
    _add_h1(doc, "Objective")
    _add_body(doc, lab["objective"])

    # Services as 2-column table
    _add_h1(doc, "AWS Services Used")
    _add_services_table(doc, lab["services"])

    # Procedure
    _add_h1(doc, "Step-by-Step Procedure")
    for step_title, lines in lab["procedure"]:
        _add_h3(doc, step_title)
        for line in lines:
            _add_bullet(doc, line)

    # Screenshots
    _add_h1(doc, "Screenshots")
    for path, caption in lab["screenshots"]:
        _add_screenshot(doc, SHOTS / path, caption)

    # Observations
    _add_h1(doc, "Observations & Results")
    _add_body(doc, lab["observations"])


# ---------------------------------------------------------------------------
# Lab content
# ---------------------------------------------------------------------------

LAB_1 = {
    "title": "Virtual Cloud Environment (VPC)",
    "objective": (
        "Create and configure a Virtual Private Cloud (VPC) on AWS in the Mumbai "
        "region. The VPC acts as the base private network within which other "
        "lab resources can later be launched in isolation from the rest of the "
        "internet."
    ),
    "services": [
        "Amazon VPC",
        "AWS Management Console",
    ],
    "procedure": [
        ("Step 1: Sign in and select region", [
            "Open https://console.aws.amazon.com and sign in with the IAM user.",
            "From the top-right region dropdown, select Asia Pacific (Mumbai) - ap-south-1.",
        ]),
        ("Step 2: Open the VPC service", [
            "Search for 'VPC' in the search bar and open the VPC dashboard.",
            "From the left navigation, click on Your VPCs.",
            "Click on the Create VPC button.",
        ]),
        ("Step 3: Configure the VPC", [
            "Select 'VPC only' as the resource to create.",
            "Set Name tag to vpc-mukesh.",
            "Set IPv4 CIDR block to 10.20.0.0/16.",
            "Leave IPv6 CIDR block as 'No IPv6 CIDR block' and Tenancy as Default.",
            "Click Create VPC.",
        ]),
        ("Step 4: Verify creation", [
            "Confirm the success banner appears with the new VPC ID.",
            "Verify that the State is Available and DNS resolution is Enabled.",
            "Note down the VPC ID for reference.",
        ]),
    ],
    "screenshots": [
        ("lab-01/1.1-vpc-list.png",
         "Figure 1.1 — VPC vpc-mukesh listed in the Mumbai region with State = Available."),
        ("lab-01/1.2-vpc-details.png",
         "Figure 1.2 — Details panel showing CIDR 10.20.0.0/16 and DNS settings enabled."),
    ],
    "observations": (
        "The VPC vpc-mukesh was created successfully in the ap-south-1 region with "
        "the IPv4 CIDR block 10.20.0.0/16. DNS resolution and DNS hostnames were "
        "enabled by default. The VPC entered the Available state immediately after "
        "creation, confirming that the base network for the remaining labs has been "
        "set up correctly."
    ),
}

LAB_2 = {
    "title": "Compute Instances and Startup Scripts (EC2)",
    "objective": (
        "Launch an Amazon EC2 instance and use a startup script (user_data) to "
        "automatically install the Apache web server and serve a custom welcome "
        "page on port 80 the moment the instance boots."
    ),
    "services": [
        "Amazon EC2",
        "Amazon Machine Image (Amazon Linux 2023)",
        "Security Groups",
        "AWS Management Console",
    ],
    "procedure": [
        ("Step 1: Sign in and open EC2 console", [
            "Sign in to the AWS Console and select region ap-south-1 (Mumbai).",
            "Open the EC2 service from the search bar and click Launch instance.",
        ]),
        ("Step 2: Configure the instance", [
            "Set Name to ec2-mukesh.",
            "Choose AMI: Amazon Linux 2023 (64-bit x86, free tier eligible).",
            "Select Instance type: t2.micro (free tier eligible).",
            "Skip the key pair section (or proceed without one for this lab).",
        ]),
        ("Step 3: Configure security group", [
            "Create a new security group named mukesh-ec2-sg.",
            "Add inbound rule: SSH on port 22, source 0.0.0.0/0.",
            "Add inbound rule: HTTP on port 80, source 0.0.0.0/0.",
            "Leave outbound as default (all traffic allowed).",
        ]),
        ("Step 4: Provide user data startup script", [
            "Expand the Advanced details section.",
            "In the User data field paste the bash script that installs httpd and "
            "writes the custom index.html (welcome page mentioning Mukesh, Roll No. 29).",
            "Click Launch instance.",
        ]),
        ("Step 5: Verify and test", [
            "After about 60 to 90 seconds, refresh the Instances page until the "
            "instance shows State = Running and 2/2 status checks passed.",
            "Copy the Public IPv4 address.",
            "Open http://<public-ip> in a browser to confirm Apache is serving "
            "the welcome page.",
        ]),
    ],
    "screenshots": [
        ("lab-02/2.1-ec2-running.png",
         "Figure 2.1 — EC2 instance ec2-mukesh in the Running state with status checks passed."),
        ("lab-02/2.2-ec2-details.png",
         "Figure 2.2 — Instance details: t2.micro, Amazon Linux 2023, attached security group mukesh-ec2-sg."),
        ("lab-02/2.3-browser-welcome.png",
         "Figure 2.3 — Browser at the public IP showing the custom welcome page served by Apache."),
    ],
    "observations": (
        "The EC2 instance ec2-mukesh booted successfully in the ap-south-1 region. "
        "The user_data script ran during the first boot and installed Apache (httpd), "
        "started the service, and wrote a custom index.html. Once the instance "
        "reached the Running state, the public IP responded over HTTP with the "
        "welcome page, which confirmed that the startup script executed correctly "
        "and the security group allowed inbound HTTP traffic."
    ),
}

LAB_3 = {
    "title": "Object Storage and Static Website Hosting (S3)",
    "objective": (
        "Create an Amazon S3 bucket and configure it for static website hosting. "
        "Upload an HTML page and access the rendered site through the S3 website "
        "endpoint without using any web server."
    ),
    "services": [
        "Amazon S3",
        "S3 Static Website Hosting",
        "S3 Bucket Policy",
        "AWS Management Console",
    ],
    "procedure": [
        ("Step 1: Create the bucket", [
            "Open the S3 service in the AWS Console.",
            "Click Create bucket.",
            "Enter a globally unique bucket name (mukesh-static-<random>) and "
            "choose AWS Region: Asia Pacific (Mumbai) ap-south-1.",
        ]),
        ("Step 2: Allow public access", [
            "In Block Public Access settings, uncheck 'Block all public access'.",
            "Acknowledge the warning by ticking the checkbox.",
            "Click Create bucket.",
        ]),
        ("Step 3: Enable static website hosting", [
            "Open the bucket and go to the Properties tab.",
            "Scroll to Static website hosting and click Edit.",
            "Choose 'Enable', set Hosting type to 'Host a static website'.",
            "Set Index document to index.html and Error document to error.html.",
            "Save changes.",
        ]),
        ("Step 4: Set bucket policy and upload files", [
            "Go to the Permissions tab and add a bucket policy that allows "
            "s3:GetObject for everyone (Principal *).",
            "Go to the Objects tab and upload index.html and error.html.",
        ]),
        ("Step 5: Access the website", [
            "Return to Properties → Static website hosting and copy the bucket "
            "website endpoint URL.",
            "Open the endpoint in a browser to view the rendered page.",
        ]),
    ],
    "screenshots": [
        ("lab-03/3.1-bucket-overview.png",
         "Figure 3.1 — S3 bucket created in the Mumbai region with the unique mukesh-static name."),
        ("lab-03/3.2-static-hosting.png",
         "Figure 3.2 — Static website hosting enabled in bucket properties with the website endpoint."),
        ("lab-03/3.3-browser-site.png",
         "Figure 3.3 — The static website rendered in a browser via the S3 website endpoint."),
    ],
    "observations": (
        "The S3 bucket was created with a globally unique name and configured to "
        "host a static website. After disabling Block Public Access and adding "
        "a bucket policy that allows public read on objects, the index.html and "
        "error.html files were uploaded. The bucket website endpoint served the "
        "page over HTTP and rendered correctly in the browser, demonstrating "
        "that S3 alone can host a fully static site without any backend server."
    ),
}

LAB_4 = {
    "title": "Virtual Networking — Subnets, Routing, and Security Groups",
    "objective": (
        "Build a custom VPC with one public subnet and one private subnet across "
        "two availability zones, attach an internet gateway, configure route "
        "tables, and define two tiered security groups (web and database) to "
        "demonstrate a typical multi-tier network design."
    ),
    "services": [
        "Amazon VPC",
        "Subnets (Public and Private)",
        "Internet Gateway",
        "Route Tables",
        "Security Groups",
    ],
    "procedure": [
        ("Step 1: Create the custom VPC", [
            "Open the VPC dashboard.",
            "Create a new VPC named vpc-mukesh-net with CIDR 10.30.0.0/16.",
            "Enable DNS hostnames and DNS resolution.",
        ]),
        ("Step 2: Create subnets", [
            "Create a public subnet subnet-mukesh-public with CIDR 10.30.1.0/24 "
            "in availability zone ap-south-1a.",
            "Enable Auto-assign public IPv4 address on the public subnet.",
            "Create a private subnet subnet-mukesh-private with CIDR 10.30.2.0/24 "
            "in availability zone ap-south-1b.",
        ]),
        ("Step 3: Attach an internet gateway", [
            "Create an Internet Gateway named igw-mukesh.",
            "Attach the IGW to vpc-mukesh-net.",
        ]),
        ("Step 4: Configure route tables", [
            "Create a public route table rt-mukesh-public in vpc-mukesh-net.",
            "Add a route 0.0.0.0/0 -> igw-mukesh.",
            "Associate rt-mukesh-public with subnet-mukesh-public.",
            "Create a private route table rt-mukesh-private and associate it with "
            "subnet-mukesh-private (no internet route).",
        ]),
        ("Step 5: Create security groups", [
            "Create mukesh-web-sg in vpc-mukesh-net allowing inbound HTTP (80) and "
            "HTTPS (443) from 0.0.0.0/0.",
            "Create mukesh-db-sg in vpc-mukesh-net allowing inbound MySQL (3306) "
            "only from the source mukesh-web-sg, not from the open internet.",
        ]),
    ],
    "screenshots": [
        ("lab-04/4.1-vpc-overview.png",
         "Figure 4.1 — Custom VPC vpc-mukesh-net details with CIDR 10.30.0.0/16."),
        ("lab-04/4.2-subnets.png",
         "Figure 4.2 — Public and private subnets in two availability zones."),
        ("lab-04/4.3-route-table-public.png",
         "Figure 4.3 — Public route table with the 0.0.0.0/0 route pointing to the internet gateway."),
        ("lab-04/4.4-igw-attached.png",
         "Figure 4.4 — Internet gateway igw-mukesh attached to vpc-mukesh-net."),
        ("lab-04/4.5-sg-web-rules.png",
         "Figure 4.5 — mukesh-web-sg inbound rules allowing HTTP and HTTPS from anywhere."),
        ("lab-04/4.6-sg-db-rules.png",
         "Figure 4.6 — mukesh-db-sg allowing MySQL only from the mukesh-web-sg source."),
    ],
    "observations": (
        "A custom multi-tier VPC was successfully built. The public subnet is "
        "internet-facing through the IGW and its associated public route table, "
        "while the private subnet has no direct internet route, which is the "
        "expected behaviour for backend or database tiers. The two security "
        "groups model a typical web-to-database trust boundary - the web tier "
        "is reachable from the public internet, but the database tier accepts "
        "MySQL traffic only from the web SG. This setup forms the standard "
        "starting point for any production-style AWS workload."
    ),
}

LAB_5 = {
    "title": "Load Balancer and Auto-Scaling Simulation",
    "objective": (
        "Configure an Application Load Balancer (ALB) in front of an Auto Scaling "
        "Group (ASG) of EC2 instances spread across two availability zones, "
        "demonstrating horizontal scaling and even distribution of incoming HTTP "
        "traffic."
    ),
    "services": [
        "Application Load Balancer (ALB)",
        "Target Groups",
        "Launch Templates",
        "Auto Scaling Groups",
        "Amazon EC2",
    ],
    "procedure": [
        ("Step 1: Build the supporting VPC", [
            "Create a VPC vpc-mukesh-alb with CIDR 10.40.0.0/16.",
            "Create two public subnets in ap-south-1a and ap-south-1b with auto "
            "public IP enabled.",
            "Attach an internet gateway and add the default route to the route table.",
        ]),
        ("Step 2: Create launch template", [
            "Create launch template lt-mukesh using Amazon Linux 2023 and t2.micro.",
            "Provide a user-data script that installs Apache and writes an index.html "
            "showing the EC2 instance ID and availability zone.",
            "Attach the security group mukesh-alb-web-sg (inbound HTTP from anywhere).",
        ]),
        ("Step 3: Create target group and ALB", [
            "Create target group tg-mukesh: HTTP, port 80, target type Instance, "
            "with health check path '/'.",
            "Create an internet-facing Application Load Balancer alb-mukesh in "
            "the two public subnets.",
            "Add an HTTP:80 listener that forwards to tg-mukesh.",
        ]),
        ("Step 4: Create the Auto Scaling Group", [
            "Create asg-mukesh referencing the launch template above.",
            "Set min size = 2, desired = 2, max size = 3.",
            "Attach to the target group tg-mukesh and use VPC zone identifier "
            "with both public subnets.",
            "Set health check type to ELB.",
        ]),
        ("Step 5: Verify load balancing", [
            "Wait for the two ASG instances to pass target group health checks.",
            "Open the ALB DNS name in a browser.",
            "Refresh the page multiple times - each refresh should sometimes "
            "return a different instance ID, proving that the ALB distributes "
            "requests across both backend instances.",
        ]),
    ],
    "screenshots": [
        ("lab-05/5.1-alb-overview.png",
         "Figure 5.1 — Application Load Balancer alb-mukesh with internet-facing scheme and DNS name."),
        ("lab-05/5.2-asg-instances.png",
         "Figure 5.2 — Auto Scaling Group asg-mukesh with two InService instances."),
        ("lab-05/5.3-target-group-health.png",
         "Figure 5.3 — Target group tg-mukesh showing both targets healthy."),
        ("lab-05/5.4-browser-alb-1.png",
         "Figure 5.4 — First request via the ALB DNS, served by the first ASG instance."),
        ("lab-05/5.5-browser-alb-2.png",
         "Figure 5.5 — After a refresh, the request is served by the second ASG instance, confirming load balancing."),
    ],
    "observations": (
        "The ALB and ASG combination worked as expected. Both instances launched "
        "by the ASG passed the target-group health check within roughly two "
        "minutes of boot, after which the ALB started routing traffic to them. "
        "Hitting the ALB DNS name and refreshing the page returned different "
        "instance IDs across requests, which confirmed that the load balancer "
        "was distributing traffic across both backend hosts. The ASG also "
        "demonstrates self-healing - if either instance is terminated, the ASG "
        "would launch a replacement to maintain the desired capacity of two."
    ),
}

LAB_6 = {
    "title": "IAM Users, Groups and Policy Configuration",
    "objective": (
        "Create IAM users, an IAM group, and attach AWS managed policies to "
        "demonstrate role-based access control - one developer user with EC2 "
        "permissions through group membership, and one read-only user with a "
        "policy attached directly."
    ),
    "services": [
        "AWS Identity and Access Management (IAM)",
        "IAM Users",
        "IAM Groups",
        "AWS Managed Policies",
    ],
    "procedure": [
        ("Step 1: Create the IAM group", [
            "Open the IAM console.",
            "Go to User groups and click Create group.",
            "Set group name to Developers-mukesh.",
            "Attach the AWS managed policy AmazonEC2FullAccess to the group.",
        ]),
        ("Step 2: Create the developer user", [
            "Go to Users and click Create user.",
            "Set user name to mukesh-dev.",
            "Enable AWS Management Console access and let the system generate a password.",
            "Untick 'Require password reset on next sign-in' for this lab.",
            "Add the user to the Developers-mukesh group on the next screen.",
            "Complete the wizard and note down the auto-generated password.",
        ]),
        ("Step 3: Create the read-only user", [
            "Create another user named mukesh-readonly.",
            "On the Set permissions screen choose 'Attach policies directly'.",
            "Search for and attach the AWS managed policy ReadOnlyAccess.",
            "Complete the wizard.",
        ]),
        ("Step 4: Sign in as the IAM user", [
            "Open an Incognito/private browser window.",
            "Go to the account-specific console URL: "
            "https://<account-id>.signin.aws.amazon.com/console.",
            "Sign in with the IAM user name mukesh-dev and its password.",
            "Verify the top-right banner shows mukesh-dev signed in.",
        ]),
    ],
    "screenshots": [
        ("lab-06/6.1-iam-users-list.png",
         "Figure 6.1 — IAM Users page showing both mukesh-dev and mukesh-readonly."),
        ("lab-06/6.2-iam-group-developers.png",
         "Figure 6.2 — The Developers-mukesh group with AmazonEC2FullAccess attached."),
        ("lab-06/6.3-iam-user-readonly-policies.png",
         "Figure 6.3 — mukesh-readonly user with the ReadOnlyAccess policy attached directly."),
        ("lab-06/6.4-iam-login-as-dev.png",
         "Figure 6.4 — Successful sign-in as mukesh-dev confirmed by the top-right banner."),
    ],
    "observations": (
        "Two IAM users with different access patterns were created. mukesh-dev "
        "received its EC2 permissions indirectly, by being a member of the "
        "Developers-mukesh group which had AmazonEC2FullAccess attached. "
        "mukesh-readonly received its permissions directly through a policy "
        "attached at the user level. The sign-in test in an incognito window "
        "succeeded with the auto-generated password, confirming that the login "
        "profile was created correctly. Both styles - group-based and "
        "user-attached - are valid, but in real environments the group-based "
        "pattern scales much better because permission changes apply to all "
        "members at once."
    ),
}

LAB_7 = {
    "title": "Serverless Function Deployment (Lambda, Fargate, DynamoDB)",
    "objective": (
        "Deploy and test three serverless / managed services: an AWS Lambda "
        "function that writes records to DynamoDB on each invocation, the "
        "DynamoDB table itself, and an ECS Fargate task running an Nginx "
        "container - all without managing any underlying servers."
    ),
    "services": [
        "AWS Lambda",
        "Amazon DynamoDB",
        "Amazon ECS with the Fargate launch type",
        "AWS IAM (execution roles)",
        "Amazon CloudWatch Logs",
    ],
    "procedure": [
        ("Step 1: Create the DynamoDB table", [
            "Open the DynamoDB console.",
            "Create a table named visitors-mukesh.",
            "Set the partition key to id (String).",
            "Choose On-demand (PAY_PER_REQUEST) capacity mode.",
        ]),
        ("Step 2: Create IAM role and Lambda function", [
            "Create an IAM role role-mukesh-lambda with a trust policy for "
            "lambda.amazonaws.com.",
            "Attach AWSLambdaBasicExecutionRole and a custom inline policy "
            "allowing dynamodb:PutItem on the visitors-mukesh table.",
            "Open the Lambda console and create a function named "
            "lambda-mukesh-visitor-logger using Python 3.12.",
            "Use role-mukesh-lambda as the execution role.",
            "Set environment variable TABLE_NAME = visitors-mukesh.",
            "Paste in the handler.py code that builds an item and calls "
            "ddb.put_item().",
        ]),
        ("Step 3: Test the Lambda function", [
            "From the Lambda console click the Test tab.",
            "Configure a test event with body {} and save it.",
            "Click Test - the response should be a 200 with a JSON body that "
            "contains the new id and the table name it wrote to.",
            "Click Test two more times so that DynamoDB ends up with three rows.",
        ]),
        ("Step 4: Verify items in DynamoDB", [
            "Open DynamoDB → Tables → visitors-mukesh.",
            "Click Explore table items. Three items should appear, each with a "
            "unique UUID id, a numeric timestamp, name = Mukesh and source = "
            "lambda-mukesh-visitor-logger.",
        ]),
        ("Step 5: Run a Fargate task", [
            "Create an ECS cluster fargate-mukesh-cluster.",
            "Create a task definition nginx-mukesh: Fargate, 0.25 vCPU, 512 MB, "
            "container image public.ecr.aws/nginx/nginx:alpine.",
            "Run a one-off task on the cluster, in the default VPC subnets, with "
            "Auto-assign public IP turned on.",
            "Wait until the task moves to Last status = RUNNING and view the "
            "CloudWatch logs to confirm the Nginx container started.",
            "Stop the task manually after capturing the screenshots.",
        ]),
    ],
    "screenshots": [
        ("lab-07/7.1-lambda-function.png",
         "Figure 7.1 — Lambda function lambda-mukesh-visitor-logger with Python 3.12 runtime."),
        ("lab-07/7.2-lambda-test-success.png",
         "Figure 7.2 — Lambda Test tab showing a successful invocation with the response body."),
        ("lab-07/7.3-dynamodb-items.png",
         "Figure 7.3 — DynamoDB visitors-mukesh table with three items written by the Lambda."),
        ("lab-07/7.4-fargate-task-running.png",
         "Figure 7.4 — Fargate task in the RUNNING state on the fargate-mukesh-cluster."),
        ("lab-07/7.5-fargate-task-logs.png",
         "Figure 7.5 — CloudWatch logs from the Fargate task showing Nginx startup messages."),
    ],
    "observations": (
        "All three serverless services worked end to end. The Lambda function, "
        "running on the AWS-managed runtime, successfully assumed its execution "
        "role, wrote items into DynamoDB on each invocation, and returned a 200 "
        "response. DynamoDB stored every record reliably with no provisioning "
        "required, since the table was created in on-demand mode. The ECS "
        "Fargate task pulled the public Nginx image from ECR Public, started "
        "the container without any EC2 host being managed by the user, and "
        "streamed its stdout into CloudWatch Logs - which is the typical "
        "developer experience for containerised serverless workloads on AWS."
    ),
}

LAB_8 = {
    "title": "Messaging Queue and Pub/Sub Simulation (SNS, SQS)",
    "objective": (
        "Configure an SNS topic with two subscribers - an SQS queue and an email "
        "address - to demonstrate the publish/subscribe (fan-out) messaging "
        "pattern, where a single published message is delivered to every "
        "confirmed subscriber."
    ),
    "services": [
        "Amazon SNS (Simple Notification Service)",
        "Amazon SQS (Simple Queue Service)",
        "Email subscription protocol",
        "SQS access policy (queue policy)",
    ],
    "procedure": [
        ("Step 1: Create the SNS topic", [
            "Open the SNS console.",
            "Create a Standard topic named sns-mukesh-notifications.",
            "Leave the rest of the settings as default.",
        ]),
        ("Step 2: Create the SQS queue", [
            "Open the SQS console.",
            "Create a Standard queue named sqs-mukesh-orders.",
            "Apply a queue access policy that allows the SNS topic to call "
            "sqs:SendMessage on this queue (using SourceArn = topic ARN).",
        ]),
        ("Step 3: Subscribe the SQS queue to the topic", [
            "From the SNS topic page, click Create subscription.",
            "Protocol: Amazon SQS, Endpoint: ARN of sqs-mukesh-orders.",
            "Click Create. The subscription becomes Confirmed automatically.",
        ]),
        ("Step 4: Subscribe an email address", [
            "Create a second subscription on the same topic.",
            "Protocol: Email, Endpoint: a real inbox address.",
            "Wait for the AWS Notification - Subscription Confirmation email.",
            "Open it and click the Confirm subscription link.",
            "The subscription's Status changes to Confirmed.",
        ]),
        ("Step 5: Publish and verify fan-out", [
            "Back on the topic, click Publish message.",
            "Set a Subject (e.g., 'Test from Mukesh - Lab 8') and a Message body.",
            "Click Publish.",
            "Within a few seconds the email subscriber receives the message.",
            "In the SQS console, open sqs-mukesh-orders and click Poll for messages. "
            "The same message arrives wrapped in an SNS notification envelope.",
        ]),
    ],
    "screenshots": [
        ("lab-08/8.1-sns-topic.png",
         "Figure 8.1 — SNS topic sns-mukesh-notifications with two confirmed subscriptions."),
        ("lab-08/8.2-email-received.png",
         "Figure 8.2 — The published message delivered to the email subscriber."),
        ("lab-08/8.3-sqs-poll-result.png",
         "Figure 8.3 — The SQS queue receiving the same message wrapped in an SNS envelope."),
    ],
    "observations": (
        "The pub/sub fan-out pattern worked exactly as expected. A single Publish "
        "from the SNS topic was delivered both to the email subscriber and to "
        "the SQS queue - the SQS message arrived inside the standard SNS JSON "
        "envelope that includes Type, MessageId, TopicArn, Subject and the "
        "original Message field. This is the canonical decoupling pattern used "
        "in real AWS systems: a producer publishes once, and many independent "
        "consumers (queues, email, HTTPS endpoints, Lambda, mobile push) can "
        "process the event in their own time, without the producer needing to "
        "know about any of them."
    ),
}

LAB_9 = {
    "title": "Infrastructure as Code using CloudFormation",
    "objective": (
        "Provision AWS infrastructure declaratively using AWS CloudFormation. "
        "A YAML template describes the desired resources (a VPC and an S3 "
        "bucket); CloudFormation then handles their creation, ordering and "
        "tracking as a single managed unit called a stack."
    ),
    "services": [
        "AWS CloudFormation",
        "Amazon VPC (managed by CloudFormation)",
        "Amazon S3 (managed by CloudFormation)",
        "AWS CLI",
    ],
    "procedure": [
        ("Step 1: Author the template", [
            "Create a YAML file stack-mukesh.yaml.",
            "Add the AWSTemplateFormatVersion and a Description.",
            "Define two Parameters - OwnerName (default 'Mukesh') and VpcCidr "
            "(default 10.50.0.0/16).",
            "Define two Resources - an AWS::EC2::VPC named CfnVpc with CIDR "
            "from the parameter, and an AWS::S3::Bucket named CfnBucket using "
            "!Sub with the AWS::AccountId pseudo-parameter for uniqueness.",
            "Define four Outputs - VpcId, VpcCidrOutput, BucketName and BucketArn.",
        ]),
        ("Step 2: Deploy the stack", [
            "Run aws cloudformation deploy --template-file stack-mukesh.yaml "
            "--stack-name stack-mukesh --region ap-south-1 --capabilities "
            "CAPABILITY_NAMED_IAM.",
            "Wait for the message Successfully created/updated stack.",
            "Alternatively, the same template can be uploaded through the "
            "CloudFormation console UI.",
        ]),
        ("Step 3: Inspect the stack in the console", [
            "Open the CloudFormation console for region ap-south-1.",
            "Locate stack-mukesh in the Stacks list (Status = CREATE_COMPLETE).",
            "Open the stack and review the Events tab for the lifecycle of "
            "each resource.",
            "Open the Resources tab to see the two logical IDs (CfnVpc, CfnBucket) "
            "mapped to their actual physical IDs.",
            "Open the Outputs tab to see the values returned by the stack.",
        ]),
        ("Step 4: Delete the stack", [
            "Run aws cloudformation delete-stack --stack-name stack-mukesh.",
            "Wait until the stack disappears from the console - both the VPC "
            "and the S3 bucket are deleted by CloudFormation in the correct order.",
        ]),
    ],
    "screenshots": [
        ("lab-09/9.1-stack-list.png",
         "Figure 9.1 — CloudFormation stack stack-mukesh in CREATE_COMPLETE state."),
        ("lab-09/9.2-stack-events.png",
         "Figure 9.2 — Stack events showing CREATE_IN_PROGRESS to CREATE_COMPLETE for each resource."),
        ("lab-09/9.3-stack-resources.png",
         "Figure 9.3 — Stack resources panel: logical IDs mapped to physical resource IDs."),
        ("lab-09/9.4-stack-outputs.png",
         "Figure 9.4 — Stack outputs: VpcId, VpcCidrOutput, BucketName and BucketArn."),
    ],
    "observations": (
        "CloudFormation deployed the entire stack from a single YAML file. The "
        "Events tab confirmed that resources were created in the correct order "
        "and the stack reached CREATE_COMPLETE without manual intervention. The "
        "Outputs tab exposed values such as the VPC ID and the bucket name, "
        "which other stacks (or operators) can consume. Compared with the "
        "earlier labs that used Terraform, CloudFormation offers very similar "
        "Infrastructure as Code benefits - the template is version-controllable, "
        "repeatable across environments and self-documenting. The main "
        "difference is that CloudFormation is AWS-native and Terraform is "
        "cross-cloud, but for a single-cloud workload either is a good choice."
    ),
}

LABS = [LAB_1, LAB_2, LAB_3, LAB_4, LAB_5, LAB_6, LAB_7, LAB_8, LAB_9]


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    configure_default_styles(doc)
    configure_sections(doc)

    add_cover_page(doc)
    add_toc(doc, LABS)

    for i, lab in enumerate(LABS, 1):
        add_lab(doc, i, lab, is_first=(i == 1))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
